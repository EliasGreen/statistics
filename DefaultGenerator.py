import random
import math
from math import sqrt

from numpy import sqrt, cos, pi
from scipy.special import erf
import scipy.integrate as integrate

import docx
from docx import Document
from docx.shared import Inches

import Combinatorics
from IGenerator import IGenerator
from Quest import Quest

from scipy.special import erf

Phi = lambda x: erf(x/2**0.5)/2

class DefaultGenerator(IGenerator):
	def __init__(self, requestQuests: list, variantNum: int, wordFileName: str):
		super().__init__(requestQuests)

		# Установка функций-генераторов для каждого номера
		self._questsFunctionsTable[1] = self.__quest1
		self._questsFunctionsTable[2] = self.__quest2
		self._questsFunctionsTable[3] = self.__quest3
		self._questsFunctionsTable[4] = self.__quest4
		self._questsFunctionsTable[5] = self.__quest5
		self._questsFunctionsTable[6] = self.__quest6
		self._questsFunctionsTable[7] = self.__quest7
		self._questsFunctionsTable[8] = self.__quest8
		self._questsFunctionsTable[9] = self.__quest9
		self._questsFunctionsTable[10] = self.__quest10
		self._questsFunctionsTable[11] = self.__quest11
		self._questsFunctionsTable[12] = self.__quest12
		self._questsFunctionsTable[13] = self.__quest13
		self._questsFunctionsTable[14] = self.__quest14
		self._questsFunctionsTable[15] = self.__quest15
		self._questsFunctionsTable[16] = self.__quest16
		self._questsFunctionsTable[17] = self.__quest17
		self._questsFunctionsTable[18] = self.__quest18

		# Настройка документа. Добавление стилей.
		self._document.styles['Normal'].font.name = 'Times New Roman'
		self._document.styles['Normal'].font.size = docx.shared.Pt(14)

		self._document.styles['Header'].font.name = 'Times New Roman'
		self._document.styles['Header'].font.size = docx.shared.Pt(18)
		self._document.styles['Header'].font.bold = True

		self._document.styles.add_style('QuestHeader', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
		self._document.styles['QuestHeader'].font.name = 'Times New Roman'
		self._document.styles['QuestHeader'].font.size = docx.shared.Pt(16)
		self._document.styles['QuestHeader'].font.bold = True

		header = self._document.add_heading(f'Типовой расчет. Вариант {variantNum}\n', 0)
		header.style = self._document.styles['Header']
		header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

		counter = 1
		for number in self._request:
			try:
				self._questsFunctionsTable[number]
				self._quests[counter] = self._questsFunctionsTable[number]()
				counter += 1

			except KeyError as exp:
				print("Чо за номер, ну шо такое")

		# Запись данных в вордовский документ
		# for number, quest in self._quests.items():
		# 	header = self._document.add_heading(f"\tЗадание {number}\n", 2)
		# 	header.style = self._document.styles['QuestHeader']

		# 	paragraph = self._document.add_paragraph(f"\t{quest._wording}\n", style = self._document.styles['Normal'])
		# 	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		# 	header = self._document.add_heading(f"\tЗадание {number}, ХОД РЕШЕНИЯ\n", 2)
		# 	header.style = self._document.styles['QuestHeader']

		# 	paragraph = self._document.add_paragraph(f"\t{quest._decisionProgress}\n\tОтвет: {quest.answer}\n", style = self._document.styles['Normal'])
		# 	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		# print(self._quests[7]._decisionProgress)
		self._document.save(wordFileName + '.docx')

	# Задача про шары. В корзине некоторое количество шаров. Вытаскивают кучку. Вытащенные шары ОДИНАКОВЫ
	def __quest1(self) -> Quest:
		white = random.randint(5, 15)
		black = random.randint(5, 15)
		n = white + black
		pick = random.randint(2, int(min(black, white) / 2))

		templatesList = [
			f"В корзине размещено {white} белых шаров и {black} черных шаров. Из них наудачу вынимают {pick} шаров. Найти"
			f" вероятность того, что они одного цвета.",

			f"В шляпе находится {white} красных фантов и {black} желтых фантов. Игрок наудачу вынимает {pick} фантов. Найти"
			f" вероятность того, что они одного цвета.",

			f"В шкафу лежит {white + black} носков, из которых {white} белых, а остальные черные. Петя наудачу берет {pick}"
			f" носков из шкафа. Какова вероятность того, что они одного цвета?"
		]

		template = random.choice(templatesList)
		resultsAll = Combinatorics.combinations(n, pick)
		resultsWhite = Combinatorics.combinations(white, pick)
		resultsBlack = Combinatorics.combinations(black, pick)

		answer = float(f"{(resultsWhite + resultsBlack) / resultsAll : .4f}") # Эта как - единственное, шо я нагуглил

		decisionProgress =  f"Количество исходов в данной задаче равняется количеству способов выбрать {pick} предметов из {n}."
		decisionProgress += f" Представим событие А - выбрано {pick} одинаковых предмета - как сумму двух событий A1 и A2."
		decisionProgress += f" Теперь выберем необходимые предметы из своих кучек и разделим количество способов выбора на все исходы."
		decisionProgress += f" После, просуммируем вероятности несовместных событий А1 и А2:\n"
		decisionProgress += f"\tP = ({resultsWhite} * {resultsBlack}) / {resultsAll} = {answer : .4f}.\n"

		header = self._document.add_heading(f"\tЗадание 1\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 1, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Задача про шары. В корзине некоторое количество шаров. Вытаскивают кучку. Вытащенные шары не являются одинаковыми
	def __quest2(self) -> Quest:
		white = 0
		black = 0

		while (white == black):
			white = random.randint(5, 15)
			black = random.randint(5, 15)

		whitePick = random.randint(2, int(white / 2.3))
		blackPick = random.randint(2, int(black / 2.3))
		n = white + black
		pickNumber = whitePick + blackPick

		templatesList = [
			f"В корзине размещено {white} белых шаров и {black} черных шаров. Из них наудачу вынимают {pickNumber} шаров. Найти"
			f" вероятность того, что будет вынуто ровно {whitePick} белых и {blackPick} черных шаров.",

			f"В шляпе находится {white} красных фантов и {black} желтых фантов. Игрок наудачу вынимает {pickNumber} фантов. Найти"
			f" вероятность того, что будет вынуто ровно {whitePick} красных и {blackPick} желтых фантов.",

			f"В шкафу лежит {white + black} носков, из которых {white} белых, а остальные черные. Петя наудачу берет {pickNumber}"
			f" носков из шкафа. Какова вероятность того, что будет вынуто ровно {whitePick} белых и {blackPick}"
			f" черных носков?"
		]

		template = random.choice(templatesList)
		resultsAll = Combinatorics.combinations(n, pickNumber)
		resultsWhite = Combinatorics.combinations(white, whitePick)
		resultsBlack = Combinatorics.combinations(black, blackPick)

		answer = float(f"{resultsWhite * resultsBlack / resultsAll : .4f}") # Эта как - единственное, шо я нагуглил

		decisionProgress =  f"Количество исходов в данной задаче равняется количеству способов выбрать {pickNumber} предметов из {n}. "
		decisionProgress += f"А это, в свою очередь является количеством сочитаний C из {n} по {pickNumber} = {resultsAll}. Аналогичным образом "
		decisionProgress += f"вычисляем количество способов вытащить заданные предметы из общих куч. В результате получаем следующие вычисления: "
		decisionProgress += f"P = {resultsWhite} * {resultsBlack} / {resultsAll} = {answer : .4f}.\n"

		header = self._document.add_heading(f"\tЗадание 2\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 2, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Задача про карты. Вынимаем карты, аки предметы из шайтан-кейса в каэсочке и надеемся на N редких
	def __quest3(self) -> Quest:
		cardCount = 32 # Питон, где константы
		pickCards = random.randint(cardCount / 4, 3 * cardCount / 4)
		pickRareCards = random.randint(1, 3)

		cardName = random.choice(["Король", "Дама", "Туз", "Валет"])
		template =  f"В колоде {cardCount} карты. Наугад вынимают {pickCards} карт. Найти вероятность того, что среди них окажется "
		template += f"хотя бы {pickRareCards} карт достоинства \"{cardName}\"."

		decisionProgress =  f"Для верного решения задачи, необходимо последовательно выбрать от {pickRareCards} до {4} карт нужного достоинства. "
		decisionProgress += f"Выбор разного количества карт представляет из себя отдельную альтернативу. Сумма альтернаив есть "
		decisionProgress += f"ответ на задачу. Посчитаем их.\n"

		answer = 0.0
		for i in range(pickRareCards, 5):
			alternative = Combinatorics.combinations(4, i) * Combinatorics.combinations(cardCount - 4, pickCards - i) / Combinatorics.combinations(cardCount, pickCards)
			answer += alternative
			decisionProgress += f"\tМы можем выбрать {i} карт; шанс получения нужной комбинации: {alternative : .2f}.\n"

		answer = float(f"{answer : .4f}")
		decisionProgress += f"\tПроссумировав данные значения, мы получаем ответ на задачу."

		header = self._document.add_heading(f"\tЗадание 3\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 3, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest4(self) -> Quest:
		# Он тут почему-то ругался на то, что не int в randrange. Выше такого не было, мам!
		# Кроме того, я тут решил разбить вычисления на большее количество строк для наглядности упрощения
		commonCount = random.randint(8, 16)
		rareCount = random.randint(5, 10)

		commonPick = random.randint(int(commonCount / 2), int((3 / 4) * commonCount))
		rarePick = random.randint(int(rareCount / 2), int((3 / 4) * rareCount))

		# TODO Если останется время - убрать эту каку. Для удобства переписывания формулы из книги так делал
		N = commonCount + rareCount # Всего студентов
		n = rareCount # Отличники
		m = commonPick + rarePick # Наудачу отобраны
		k = rarePick  # Сколько отличников отобрано

		templatesList = [
			f"В группе {N} студентов, среди которых {n} отличников. По списку наудчу отобраны "
			f"{m} студентов. Найти вероятность того, что среди них {k} отличников.",

			f"В партии из {N} деталей, {n} стандартных. Случайно отбирают ровно "
			f"{m} деталей. Найти вероятность того, что среди них {k} стандартных."
		]
		template = random.choice(templatesList)

		allResults = Combinatorics.combinations(N, m)
		standartPick = Combinatorics.combinations(n, k)
		notStandartPick = Combinatorics.combinations(N - n, m - k)
		answer = standartPick * notStandartPick / allResults
		answer = float(f"{answer : .2f}")

		decisionProgress =  f"Количетсво возможных исходов, равное количеству способов выбрать {m} из {N} "
		decisionProgress += f"элементов, равняется {allResults}. В благоприятном исходе, необходимо выбрать {k} "
		decisionProgress += f"из выделяющийся группы в {n} элементов и {m} из {N - n} "
		decisionProgress += f"элементов из невыделяющейся группы. Итого мы имеем: {standartPick} * {notStandartPick} / {allResults} = {answer}"

		header = self._document.add_heading(f"\tЗадание 4\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 4, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Попытка не сломаться за N итераций.
	def __quest5(self) -> Quest:
		nodesCount = 4
		nodesBreakChance = [random.randint(5, 15)*0.01 for i in range(0, nodesCount)]
		nodesBreakChance = [float(f"{i : .4f}") for i in nodesBreakChance]

		# Я понимаю как хреново читается эта жопка в шансе. Но я хочу попитонить, ня!
		chancesStr = ''.join([f'P({i + 1}) = {ch}; ' for i, ch in enumerate(nodesBreakChance)])
		chancesStr = chancesStr[0 : -2]

		templatesList = [
			f"При изготовлении детали заготовка должна пройти {nodesCount} операции. Предпологая появление брака на отдельных операциях "
			f"независимыми, найти вероятность изготовления стандартной детали, если вероятность появления брака на i-й операции "
			f"P(i) равны {chancesStr}?",

			f"Узел автомашины состоит из {nodesCount} деталей. Вероятности выхода деталей из строя равны {chancesStr}"
			f" и являются независимыми друг от друга. Узел выходит из строя, если сломается хотя бы одна деталь. Какова вероятность того, "
			f"что узел автомашины не выйдет из строя?"
		]
		template = random.choice(templatesList)

		answer = 1.0
		for i in nodesBreakChance:
			answer *= 1 - i
		answer = float(f"{answer : .2f}")

		decisionProgress =  "Для решения данной задачи необходимо просто перемножить вероятности, обратные данным для каждого узла (заготовки).\n"
		decisionProgress += f"P = {''.join([f'(1 - {i})' for i in nodesBreakChance])} = {answer}."

		header = self._document.add_heading(f"\tЗадание 5\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 5, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Вытаскиваем шарики. хз куда
	# def __quest5(self) -> Quest:
	# 	# Опять не int в randrange. Шо це такое
	# 	blackCount = int()
	# 	whiteCount = int()
	# 	while blackCount == whiteCount:
	# 		blackCount = random.randint(5, 8)
	# 		whiteCount = random.randint(5, 8)
	# 	pickCount = random.randint(4, 7)

	# 	templatesList = [
	# 		f"Из урны, содержащей {blackCount} черных и {whiteCount} белых шара, вынимают "
	# 		f"{pickCount} шаров, с возвращением каждый раз вынутого шара обратно в урну. "
	# 		f"Найти вероятность того, что среди них окажется хотя бы один белый шар.",

	# 		f"В ящике {blackCount + whiteCount} деталей, из которых {whiteCount} браконванных."
	# 		f"Из ящика вынимают {pickCount} раз вынимают деталь (с возвращением ее каждый раз обратно)."
	# 		f"Найти вероятность того, что хотя бы раз будет вынута бракованная деталь."
	# 	]
	# 	template = random.choice(templatesList)

	# 	answer = (whiteCount / (blackCount + whiteCount)) ** pickCount
	# 	answer = float(f"{answer : .4f}")
	# 	decisionProgress =  f"Найдем вероятность противоположного события: {pickCount} достали черный шар (небракованную деталь): "
	# 	decisionProgress += f"({whiteCount} / {blackCount + whiteCount}) ^ {pickCount} = {answer : .4f}. Вычитая результат из единицы, "
	# 	answer = float(f"{1 - answer : .4f}")
	# 	decisionProgress += f"найдем ответ на задачу: {answer : .4f}"

	# 	header = self._document.add_heading(f"\tЗадание 5\n", 2)
	# 	header.style = self._document.styles['QuestHeader']

	# 	paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
	# 	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# 	header = self._document.add_heading(f"\tЗадание 5, ХОД РЕШЕНИЯ\n", 2)
	# 	header.style = self._document.styles['QuestHeader']

	# 	paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
	# 	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Стреляют две пушки
	def __quest6(self) -> Quest:
		first = int()
		second = int()
		while first == second:
			first = random.randint(50, 90)
			second = random.randint(40, 70)
		first = float(f"{first * 0.01 : .4f}")
		second = float(f"{second * 0.01 : .4f}")

		templatesList = [
			f"Произведен залп из двух орудий. Шанс попадания из первого орудия - {first}, а из второго - {second}."
			f"Какова вероятность поражения цели?"
		]
		template = random.choice(templatesList)

		looserLikeMe = (1 - first) * (1 - second)
		looserLikeMe = float(f"{looserLikeMe : .2f}")
		answer = float(f"{1 - looserLikeMe : .2f}")

		decisionProgress =  f"Вычислим вероятность не попасть из первого и второго орудий - {1 - first : .2f} и {1 - second : .2f} соответственно."
		decisionProgress += f"Тогда, вероятность не попасть при обоих выстрелах равна {looserLikeMe}. Получаем вероятность попадания, "
		decisionProgress += f"равной {answer}."

		header = self._document.add_heading(f"\tЗадание 6\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 6, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# Бомбы на мост. Может, тот 5-й сюда?
	def __quest7(self) -> Quest:
		bombs = random.randint(3, 4)
		bombsList = [random.randint(3, 8) * 0.1 for i in range(0, bombs)]
		bombsList = [float(f"{i : .2f}") for i in bombsList]

		templatesList = [
			f"Для разрушения моста достаточно попадания одной авиационной бомбы. Найти вероятность того, что мост будет разрушен, "
			f"если на него сбросить {bombs} бомбы, вероятности попадания которых соответственно равны "
			f"{''.join([f'{i : .2f} ' for i in bombsList])}"
		]
		template = random.choice(templatesList)

		answer = 1.0
		for i in bombsList:
			answer *= 1 - i
		answer = float(f"{answer : .2f}")

		decisionProgress =  f"Вычислим вероятность того, что ни одна бомба не попадет в цель: {''.join([f'(1 - {i : .2f})' for i in bombsList])} = {answer}"
		answer = float(f"{1 - answer : .2f}")
		decisionProgress += f"Вычитая данную вероятность из 1, получаем: {answer : .2f}"

		header = self._document.add_heading(f"\tЗадание 7\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 7, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest8(self):
		poorChance = [random.randint(2, 10) * 0.01 for i in range(0, 3)]
		poorChance = [float(f"{i : .4f}") for i in poorChance]

		template =  f"Рабочий обслуживает 3 автомата. Вероятность брака для первого равна {poorChance[0]}; для второго "
		template += f"{poorChance[1]}; для третьего {poorChance[2]}. Производительноть автоматов одинакова. Изготовленные "
		template += f"детали попадают на общий конвейер. Определить вероятность того, что взятая наугад деталь будет годной."

		answer = 0
		for i in poorChance:
			answer += 0.33 * (1 - i)
		answer = float(f"{answer : .2f}")

		decisionProgress =  f"В связи с тем, что каждый автомат имеет одинаковую производительность, шанс на то, что деталь "
		decisionProgress += f"была сделана на любом из станков одинакова и равна 1 / 3 (в дальнейшем 0.33). Теперь вычислим "
		decisionProgress += f"шанс того, что взятая наугад деталь будет годной:{(''.join([f' 0.33 * (1 - {i}) +' for i in poorChance])[0 : -1])}= {answer}"

		header = self._document.add_heading(f"\tЗадание 8\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 8, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer : .2f}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest9(self) -> Quest:
		casesCount = 3
		cases = dict() # Словарь вида Буква случая -> (Вероятность заболеть, вероятность выжить)

		def getP(max = 1) -> float:
			"""Генерирует и округляет вероятность"""
			p = random.uniform(0, max)
			return float(f"{p : .2f}")

		# Раскрою цикл так, чтобы точно оставались шансы. В цикле иногда багуется, не смог исправить (Лень))0)00
		cases['A'] = (random.randrange(10, 70), getP(0.95))
		cases['B'] = (random.randrange(10, 100 - cases['A'][0]), getP(0.95))
		cases['C'] = (100 - cases['A'][0] - cases['B'][0], getP(0.95))

		for k in cases.keys():
			cases[k] = (float(f"{cases[k][0] * 0.01 : .2f}"), cases[k][1]) # хз моно ли красивее сделать

		templatesList = [
			f"В больницу поступают в среднем {cases['A'][0]} больных с заболеванием А, {cases['B'][0]} с заболеванием В, {cases['C'][0]} с заболеванием "
			f"С. Вероятность полного выздоровления для каждого заболевания соответственно равны {cases['A'][1]}, {cases['B'][1]} и {cases['C'][1]}. "
			f"Больной был выписан здоровым. Найти вероятность того, что он страдал заболеванием А."
		]
		template = random.choice(templatesList)

		PZ = cases['A'][0] * cases['A'][1] + cases['B'][0] * cases['B'][1] + cases['C'][0] * cases['C'][1]
		PAZ = cases['A'][1] * cases['A'][0] / PZ

		PZ = float(f"{PZ : .2f}")
		answer = float(f"{PAZ : .2f}")

		decisionProgress =  f"Обозначим как P(Z) вероятность того, что больной выписан здоровым. P(Z) = {PZ}. Теперь мы можем вычислить  "
		decisionProgress += f"вероятность того, что он болел заболеванием А по формуле Бейеса: P(A / Z) = {answer}. Ответ на задачу: {answer}"

		header = self._document.add_heading(f"\tЗадание 9\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 9, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	# # TODO Нужны зависимости в генерации, а то часто бывают ответы уровня 0-0.5 и не более.
	def __quest10(self):
		n = random.randrange(12, 25)
		k = random.randrange(int(n * 0.1), int(n * 0.5))
		p = random.uniform(0.2, 0.6)
		p = float(f"{p : .2f}")

		templatesList = [
			f"За день в магазине было {n} покупателей. Вероятность найти покупку для каждого из них одинакова и равна {p}. Найти вероятность того, что {k} покупателей нашли покупку."
		]
		template = random.choice(templatesList)

		answer = Combinatorics.combinations(n, k) * p**k * (1-p)**(n-k)
		answer = float(f"{answer : .2f}")
		decisionProgress =  f"Воспользуемся распределением Бернулли и получим ответ: {answer}"

		header = self._document.add_heading(f"\tЗадание 10\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{template}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 10, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __tableGenerator(self) -> dict:
		"""Генератор таблички для 11 и 12 заданий"""
		table = dict()
		startX = random.randint(-3, 3) # Первый Х
		rand = 20 # Шаг рандома - 0.05

		# Раскрою цикл для украшения полученной таблицы
		table[startX] = random.randint(2, 6) *  0.05
		table[startX + 1] = random.randint(2, 5) * 0.05
		table[startX + 2] = random.randint(1, 4) * 0.05
		table[startX + 3] = random.randint(1, 3) * 0.05
		table[startX + 4] = 1 - table[startX] - table[startX + 1] - table[startX + 2] - table[startX + 3]

		for i in table.keys():
			table[i] = float(f"{table[i] : .4f}")

		return table

	# # TODO График
	def __quest11(self):
		table = self.__tableGenerator()

		# Генерация графика

		graphic = "FURSOV MAKE ME"
		# ...

		# Функция. (Значение - строка с границей)
		keys = [i for i in table.keys()]
		currentP = 0
		function = str()
		for i in range(0, len(keys) + 1):
			if i == 0:
				function += f"\t0 при X <= {keys[0]}\n"

			elif i < len(keys):
				currentP += table[keys[i - 1]]
				currentP = float(f"{currentP : .4f}")
				function += f"\t{currentP} при {keys[i-1]} < X <= {keys[i]}\n"

			else:
				currentP += table[keys[-1]]
				currentP = float(f"{currentP : .4f}")
				function += f"\t{currentP} при {keys[-1]} < X"

		decisionProgress = "Следуя определением многоугольника распределения и функции распределения, строим их."

		answer = "\ngraphic.png" + '\n' + function

		header = self._document.add_heading(f"\tЗадание 11\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\tСлучайная величина ε имеет распределение, заданное следующей таблицей:", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		tableDocument = self._document.add_table(rows = 2, cols = len(table))
		tableDocument.style = 'Table Grid'
		for i, (v, p) in enumerate(table.items()):
			tableDocument.cell(0, i).text = str(v)
			tableDocument.cell(1, i).text = str(p)

		paragraph = self._document.add_paragraph(f"\tПостроить многоугольник распределения и найти функцию распределения F(ε).\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 11, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest12(self):
		table = self.__tableGenerator()

		M = 0
		Mstr = str()
		for x, p in table.items():
			M += x * p
			Mstr += f"{x : .2f} * {p : .2f} + "
		M = float(f"{M : .4f}")
		Mstr += Mstr[0 : -3]

		D = 0
		Dstr = str()
		for x, p in table.items():
			D += x*x * p
			Dstr += f"{x : .2f} ^ 2 * {p : .2f} + "
		Dstr = Dstr[0 : -2] + f"{M} ^ 2"
		D -= M*M
		D = float(f"{D : .4f}")

		O = math.sqrt(D)
		O = float(f"{O : .4f}")

		header = self._document.add_heading(f"\tЗадание 12\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\tСлучайная величина ε распределение, заданное следующей таблицей:", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		tableDocument = self._document.add_table(rows = 2, cols = len(table))
		tableDocument.style = 'Table Grid'
		for i, (v, p) in enumerate(table.items()):
			tableDocument.cell(0, i).text = str(v)
			tableDocument.cell(1, i).text = str(p)

		paragraph = self._document.add_paragraph(f"\tНайти М(ε), D(ε) и σ(ε).\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		decisionProgress =   "Следуая классическим формулам производим вычисления значений, найдем требуемые функции: "
		decisionProgress += f"M = {Mstr};\n\tD = {Dstr};\n\tфункция σ(ε) равняется корню квадратнму из D(ε), следовательно она равна {O}."
		answer = "М(ε) = {M}, D(ε) = {D}, σ(ε) = {O}"

		header = self._document.add_heading(f"\tЗадание 12, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{decisionProgress}\n\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest13(self):
		task13 = ''
		solving13 = ''
		answer13 = ''

		task14 = ''
		solving14 = ''
		answer14 = ''

		# чтобы добавить вариативности 13-ому заданию - достаточно добавить новую функцию в лист
		# сначала идет функция -> потом ее собственный интеграл от нуля до x
		all_functions = [('5cos(5x)', 'sin(5x)'), ('2x', 'x^2'), ('4*x^3','x^4')]
		main_interval_points = ['PI/10', '1', '1']

		selected_item = random.randint(0, len(all_functions) - 1)

		# выводим условие задачи
		task13 += 'ε - непрерывная случайная величина с плотностью распределения p(x), заданной следующим образом:\n'
		task13 += f'\tp(x) = {all_functions[selected_item][0]}, для любого x, который принадлежит интервалу\n'
		task13 += f'\t(0;{main_interval_points[selected_item]})\n'
		task13 += f'\tp(x) = 0, для любого x, который не принадлежит интервалу\n'
		task13 += f'\t(0;{main_interval_points[selected_item]})\n'
		task13 += '\tНати функцию распределения F(x).'

		# выводим решение
		solving13 = "Используя определение функции распределеня, находим ее."

		answer13 += 'F(x) = 0, при x <= 0\n'
		answer13 += f'F(x) = {all_functions[selected_item][1]}, при 0 < x <= {main_interval_points[selected_item]}\n'
		answer13 += f'F(x) = 1, при x > {main_interval_points[selected_item]}'

		header = self._document.add_heading(f"\tЗадание 13\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{task13}", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 13, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{solving13}\tОтвет:\n{answer13}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		self.costil = selected_item

	def __quest14(self):
	# Далее идет условие/решение 14-ой задачи
		task14 = 'ε - непрерывная случайная величина примера 13.\n'
		task14 += 'Найти M(ε), D(ε), σ(ε).\n'
		M_E_values = [
			integrate.quad(lambda x: x*5*cos(5*x), 0, pi/10),
			integrate.quad(lambda x: x * (x*x), 0, 1),
			integrate.quad(lambda x: x * 4 * x * x * x, 0, 1),
					]

		xx_F_values = [
			integrate.quad(lambda x: x*x * 5 * cos(5 * x), 0, pi / 10),
			integrate.quad(lambda x: x*x * (x*x), 0, 1),
			integrate.quad(lambda x: x*x * 4 * x * x * x, 0, 1),
		]

		D_E_values = [xx_F_values[i][0] - (M_E_values[i][0]*M_E_values[i][0]) for i in range(len(M_E_values))]
		B_E_values = [sqrt(d_value) for d_value in D_E_values]

		solving14 = "Вычисляем интегралы, согласно определениям."

		answer14 = f'M(ε) = {M_E_values[self.costil][0] : .4f}\n'
		answer14 += f'\tD(ε) = {D_E_values[self.costil] : .4f}\n'
		answer14 += f'\tσ(ε) = {B_E_values[self.costil] : .4f}\n'

		header = self._document.add_heading(f"\tЗадание 14\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{task14}", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 14, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{solving14}\tОтвет:\n{answer14}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest15(self):
		task = ''
		solving = ''
		answer = ''

		n = 1
		p = 1
		q = 1
		k = 1

		while n*p*q <= 9 or not(sqrt(n*p*q).is_integer()):
			n = random.randint(100, 10000)
			p_values = [0.1, 0.2, 0.3, 0.5, 0.6, 0.7, 0.8]
			p = random.choice(p_values)
			q = (10.0 - p*10)/10.0

			k = random.randint(45, n - 50)

		task += f'Вероятность наступления события А в одном опыте равна {p}. Найти вероятность того, что событие А наступит {k} раз в {n} опытах.\n'

		solving += 'Так как n*p*q много больше 9, то воспользуемся локальной теоремой Лапласа:\n'
		solving += '\tPn(K) = Phi(x)/sqrt(n*p*q)\n'
		solving += '\tГде x = (k - np)/sqrt(n*p*q)\n'
		solving += '\tНайдем значение x:\n'
		solving += f'\tx = ({k} - {n}*{p})/sqrt({(n*p*q)}) = {(k - n*p)/sqrt(n*p*q) : .4f}\n'
		solving += '\tТ.к. Phi(x) - четная функция, то Phi(-x) = Phi(x)\n'
		solving += f'\tPhi({abs((k - n*p)/sqrt(n*p*q))}) = {Phi(abs((k - n*p)/sqrt(n*p*q))) : .4f}\n'
		solving += f'\tТогда P = {Phi(abs((k - n*p)/sqrt(n*p*q)))/sqrt(n*p*q) : .4f}\n'
		answer += f'P = {Phi(abs((k - n*p)/sqrt(n*p*q)))/sqrt(n*p*q) : .4f}\n'

		header = self._document.add_heading(f"\tЗадание 15\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{task}", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 15, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{solving}\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest16(self):
		task = ''
		solving = ''
		answer = ''
		all_a_values_list = [0.1, 0.3, 0.5, 1, 1.5, 1.8, 2, 2.2, 2.6, 3]
		a_value = random.choice(all_a_values_list)

		all_b_values_list = [0.1, 0.2, 0.3, 0.5, 0.6, 0.7, 0.8, 1]
		b_value = random.choice(all_b_values_list)

		all_low_values_list = [1, 2, 3]
		low_value = random.choice(all_low_values_list)

		all_sum_values_list = [1, 2, 3]
		high_value = low_value + random.choice(all_sum_values_list)

		task += f'Пусть ε - нормально распределенная случайная величина с параметрами a = {a_value}, σ = {b_value}. Найти P({low_value} < ε < {high_value})\n'

		solving += f'P({low_value} < ε < {high_value}) =\n'
		solving += f'\tF({high_value}) - F({low_value}) =\n'
		solving += f'\tФ(({high_value} - {a_value})/{b_value}) - Ф(({low_value} - {a_value})/{b_value}) =\n'
		solving += f'\tФ({(high_value - a_value)/b_value}) - Ф({(low_value - a_value)/b_value}) =\n'
		solving += f'\t{Phi((high_value - a_value) / b_value)} - {Phi((low_value - a_value) / b_value)} =\n'
		solving += f'\t{Phi((high_value - a_value) / b_value) - Phi((low_value - a_value) / b_value)}\n'

		answer += f'P({low_value} < ε < {high_value}) = '
		answer += f'{Phi((high_value - a_value) / b_value) - Phi((low_value - a_value) / b_value)}\n'

		header = self._document.add_heading(f"\tЗадание 16\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{task}", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 16, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{solving}\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest17(self):
		task = ''
		solving = ''
		answer = ''

		all_a_values_list = [0.1, 0.3, 0.5, 1, 1.5, 1.8, 2, 2.2, 2.6, 3]
		a_value = random.choice(all_a_values_list)

		all_b_values_list = [0.1, 0.2, 0.3, 0.5, 0.6, 0.7, 0.8, 1]
		b_value = random.choice(all_b_values_list)

		all_shift_values_list = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1]
		shift_value = random.choice(all_shift_values_list)

		task += 'Пусть ε - нормально распределенная случайная величина с параметрами a = {a_value}, б = {b_value}. Найти P(|ε - {a_value}| < {shift_value})\n'

		solving += '(Поскольку из нормально распределенной случайной величины вычитается ее математическое ожидание, можем воспользоваться следующей формулой)\n'
		solving += f'P(|ε - {a_value}| < {shift_value}) = 2Ф({shift_value}/{b_value}) =\n' + str(2*Phi(shift_value/b_value)) + '\n'

		answer += f'P(|ε - {a_value}| < {shift_value}) = '
		answer += str(2*Phi(shift_value/b_value)) + '\n'

		answer += f'P({low_value} < ε < {high_value}) = '
		answer += f'{Phi((high_value - a_value) / b_value) - Phi((low_value - a_value) / b_value)}\n'

		header = self._document.add_heading(f"\tЗадание 17\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{task}", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

		header = self._document.add_heading(f"\tЗадание 17, ХОД РЕШЕНИЯ\n", 2)
		header.style = self._document.styles['QuestHeader']

		paragraph = self._document.add_paragraph(f"\t{solving}\tОтвет: {answer}\n", style = self._document.styles['Normal'])
		paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

	def __quest18(self):
		task = ''
		solving = ''
		answer = ''
		column_values = [random.randint(-3, 3) for i in range(3)]
		row_values = [random.randint(-3, 3) for i in range(2)]
		matrix = []
		for i in range(2):
			row = []
			while len(row) < 3:
				if len(row) == 2:
					if i == 1:
						row.append(round(1 - sum(row) - sum(matrix[0]), 1))
						break
				random_num = round(random.random(), 1)
				if random_num >= 0.5:
					random_num = 0.1
				if i == 1 and sum(matrix[0]) + sum(row) + random_num <= 1:
					row.append(random_num)
				elif i == 0 and sum(row) + random_num <= 1:
					row.append(random_num)

			matrix.append(row)

		task += 'Дана таблица распределения вероятностей двумерной случайной величины (e, n):\n'
		task += "{:<5} {:<5} {:<5} {:<5}\n".format('e/n', column_values[0], column_values[1], column_values[2])
		task += "{:<5} {:<5} {:<5} {:<5}\n".format(row_values[0], matrix[0][0], matrix[0][1], matrix[0][2])
		task += "{:<5} {:<5} {:<5} {:<5}\n".format(row_values[1], matrix[1][0], matrix[1][1], matrix[1][2])
		task += 'Найти M(e), M(n), M(en), D(e), D(n), D(en)\n'

		M_e = round(row_values[0]*sum(matrix[0]) + row_values[1]*sum(matrix[1]), 2)
		solving += 'M(e) = {0}*{1} + {2}*{3} = {4}\n'.format(row_values[0], round(sum(matrix[0]), 2), row_values[1], round(sum(matrix[1]), 2), M_e)
		answer += 'M(e) = {0}\n'.format(M_e)
		M_n = round(column_values[0]*(matrix[0][0] + matrix[1][0]) + column_values[1]*(matrix[0][1] + matrix[1][1]) + column_values[2]*(matrix[0][2] + matrix[1][2]), 2)
		solving += 'M(n) = {0}*{1} + {2}*{3} + {4}*{5}= {6} \n'.format(column_values[0],
			round(matrix[0][0] + matrix[1][0], 2),
			column_values[1],
			round(matrix[0][1] + matrix[1][1], 2),
			column_values[2],
			round(matrix[0][2] + matrix[1][2], 2),
					M_n)
		answer += 'M(n) = {0}\n'.format(M_n)

		M_e_n_1 = row_values[0]*column_values[0]*matrix[0][0] + row_values[0]*column_values[1]*matrix[0][1] + row_values[0]*column_values[2]*matrix[0][2]
		M_e_n_2 = row_values[1]*column_values[0]*matrix[0][0] + row_values[1]*column_values[1]*matrix[0][1] + row_values[1]*column_values[2]*matrix[0][2]

		M_e_n = round(M_e_n_1 + M_e_n_2, 2)
		answer += 'M(en) = {0}\n'.format(M_e_n)
		solving += 'M(en) = {0}*{1}*{2} + {3}*{4}*{5} + {6}*{7}*{8} + {9}*{10}*{11} + {12}*{13}*{14} + {15}*{16}*{17} = {18}\n'.format(row_values[0],column_values[0],matrix[0][0],
			row_values[0],column_values[1],matrix[0][1],
		row_values[0],column_values[2],matrix[0][2],
		row_values[1],column_values[0],matrix[0][0],
			row_values[1],column_values[1],matrix[0][1],
		row_values[1],column_values[2],matrix[0][2],
					M_e_n)

		M_e_2 = round(row_values[0]*row_values[0]*sum(matrix[0]) + row_values[1]*row_values[1]*sum(matrix[1]), 2)
		M_n_2 = round(column_values[0]*column_values[0]*(matrix[0][0] + matrix[1][0]) + column_values[1]*column_values[1]*(matrix[0][1] + matrix[1][1]) + column_values[2]*column_values[2]*(matrix[0][2] + matrix[1][2]), 2)

		M_e_n_in_1 = pow(row_values[0]*column_values[0], 2)*matrix[0][0] + pow(row_values[0]*column_values[1], 2)*matrix[0][1] + pow(row_values[0]*column_values[2], 2)*matrix[0][2]
		M_e_n_in_2 = pow(row_values[1]*column_values[0], 2)*matrix[0][0] + pow(row_values[1]*column_values[1], 2)*matrix[0][1] + pow(row_values[1]*column_values[2], 2)*matrix[0][2]
		M_e_n_in_final = round(M_e_n_in_1 + M_e_n_in_2, 2)

		solving += 'M(e^2) = {0}\n'.format(M_e_2)
		solving += 'M(n^2) = {0}\n'.format(M_n_2)
		solving += 'M((en)^2) = {0}\n'.format(M_e_n_in_final)

		D_e = round(M_e_2 - M_e*M_e, 2)
		solving += 'D(e) = {0} - {1} = {2}\n'.format(M_e_2, round(M_e*M_e, 2), D_e)
		answer += 'D(e) = {0}\n'.format(D_e)
		D_n = round(M_n_2 - M_n*M_n, 2)
		solving += 'D(n) = {0} - {1} = {2}\n'.format(M_n_2,round( M_n*M_n, 2), D_n)
		answer += 'D(n) = {0}\n'.format(D_n)
		D_e_n = round(M_e_n_in_final - M_e_n*M_e_n, 2)
		solving += 'D(en) = {0} - {1} = {2}\n'.format(M_e_n_in_final, round(M_e_n*M_e_n, 2), D_e_n)
		answer += 'D(en) = {0}\n'.format(D_e_n)

		return task, solving, answer
