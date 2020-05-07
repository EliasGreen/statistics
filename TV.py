from tkinter import *
from tkinter import messagebox
from PIL import ImageTk, Image
import os

import docx
from docx import Document
from docx.shared import Inches

from DefaultGenerator import DefaultGenerator
from DefaultGeneratorWTanswers import DefaultGeneratorWTanswers

# --------------------------------
# Главный класс - приложение
# --------------------------------
class App(object):
    def __init__(self, master, **kwargs):
        self.master = master

        # Добавление главной картинки
        self.canvas = Canvas(root, width=256, height=188, bg="grey")
        self.canvas.place(relx=.5, rely=.3, anchor="c")
        try:
            self.img = ImageTk.PhotoImage(Image.open(R"VUwRK7oAv7M.jpg"))
            self.canvas.create_image(0, 0, anchor=NW, image=self.img)
        except Exception as exp:
            pass

        # Начальные графические параметры главного экрана приложения
        master.title("Генератор типовых расчетов по предмету 'Теория вероятности'")
        master.configure(background='gray')
        master.geometry('600x450')

        # Label кол-ва вариантов
        vars_count_label = Label(text="Введите количество вариантов для генерации:")
        vars_count_label.place(relx=.5, rely=.6, anchor="c")

        # Input кол-ва вариантов
        self.vars_count = StringVar()
        vars_count_entry = Entry(textvariable=self.vars_count)
        vars_count_entry.place(relx=.5, rely=.66, anchor="c")

        # Кнопка обработки генерации вариантов
        vars_gen_button = Button(text="Сгенерировать варианты", command=self.generate_variants)
        vars_gen_button.place(relx=.5, rely=.8, anchor="c")

        # Кнопка обработки показа вариантов
        vars_del_button = Button(text="Открыть папку с вариантами", command=self.show_docx_variants)
        vars_del_button.place(relx=.5, rely=.9, anchor="c")

    # Вспомогательная функция вывода информации об успешности генерации вариантов
    def show_gen_results(self):
        if self.vars_count.get() == '' :
            messagebox.showerror("Безуспешная генерация\n ",
                                "Вы не ввели количество вариантов для генерации")
            return False

        try:
            count = int(self.vars_count.get())
        except Exception as exp:
            messagebox.showinfo("Безуспешная генерация\n",
                                "Ввод не является числом!")
            return False

        if self.vars_count.get() != '' and count > 0:
            messagebox.showinfo("Успешная генерация\n ",
                                   "Сгенерированные варианты находятся в "
                                   "одной папке с программой.\n"
                                   "Количество сгенерированных вариантов: " + self.vars_count.get())
            return True
        elif self.vars_count.get() != '' and count <= 0:
            messagebox.showerror("Безуспешная генерация\n ",
                                "Количество вариантов для генерации "
                                "должно быть больше нуля")
            return False

    # Метод запуска генерации объектов
    def generate_variants(self):
        if self.show_gen_results():
            count = int(self.vars_count.get()) # Исключения проверены выше

            for i in range(1, count + 1):
                DG = DefaultGenerator([i for i in range(1, 19)], i, f"Типовой вариант # {i}")

                # self._document = Document()
                # self._document.styles['Normal'].font.name = 'Times New Roman'
                # self._document.styles['Normal'].font.size = docx.shared.Pt(14)

                # self._document.styles['Header'].font.name = 'Times New Roman'
                # self._document.styles['Header'].font.size = docx.shared.Pt(18)
                # self._document.styles['Header'].font.bold = True

                # self._document.styles.add_style('QuestHeader', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                # self._document.styles['QuestHeader'].font.name = 'Times New Roman'
                # self._document.styles['QuestHeader'].font.size = docx.shared.Pt(16)
                # self._document.styles['QuestHeader'].font.bold = True

                # header = self._document.add_heading(f'Типовой расчет. Вариант {i}\n', 0)
                # header.style = self._document.styles['Header']
                # header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                # for idx in range(len(DG.tasks)):
                #     header = self._document.add_heading(f"\tЗадание {idx+1}\n", 2)
                #     header.style = self._document.styles['QuestHeader']

                #     paragraph = self._document.add_paragraph(f"\t{DG.tasks[idx]}\n", style=self._document.styles['Normal'])
                #     paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                # self._document.save(f'Типовой вариант # {i}.docx')

    # Метод открытия папки с вариантами
    def show_docx_variants(self):
        os.system('explorer ' + os.path.dirname(os.path.abspath(__file__)))

def on_exit():
    root.destroy
    sys.exit()

# pyinstaller Combinatorics.py DefaultGenerator.py DefaultGeneratorWTanswers.py IGenerator.py Quest.py TV.py --noconsole --onefile
if __name__ == "__main__":
    root = Tk()
    app = App(root)
    try:
        root.wm_iconbitmap(R"favicon.ico")
    except Exception as exp:
        pass
    root.resizable(width=False, height=False)
    root.protocol("WM_DELETE_WINDOW", on_exit)
    root.mainloop()
